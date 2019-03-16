import nltk
import os
from nltk import RegexpTokenizer, Counter
from nltk.corpus import wordnet
import re
from nltk.stem.porter import PorterStemmer
from nltk.stem.lancaster import LancasterStemmer
from nltk.stem import SnowballStemmer
from nltk.stem import WordNetLemmatizer
'''
from docx import Document

document = Document()
document.save('mn.docx')

import docx

doc = docx.Document('mn.docx')
fullText = []
for para in doc.paragraphs:
    fullText.append(para.text)
print(fullText)'''
file=open('C:\/Users\Mariam\Desktop\data.txt','r')
f = file.read()
token2 = nltk.sent_tokenize(f)
print(token2)
token3=[]
for i in token2:
    token = nltk.word_tokenize(i)
    token = re.split('[-_,.\s ]', f)
    token = re.split('/d',f)
    token3.append(token)
print("mariam")
print (token3)
words = re.split('[-_,.\s ]', f)
print(words)
count=[]
for w in words:
    count.append(words.count(w))
print(count)

porter_stemmer = PorterStemmer()
stems=[]
for w in range(len(token3)):
    for j in range(len(token3[w])):
        stem = porter_stemmer.stem(token3[w][j])
        stems.append(stem)
print(stems)
count1=[]
for w in stems:
    count1.append(stems.count(w))
print(count1)


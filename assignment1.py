from string import punctuation
from collections import Counter
import nltk
import re


#convert from docx to txt
''''
from docx import Document

document = Document()
document.save('mn.docx')

import docx

doc = docx.Document('mn.docx')
fullText = []
for para in doc.paragraphs:
    fullText.append(para.text)
print(fullText)'''

file=open('C:\/Users\Mariam\Desktop\data.txt','r')
f = file.read()
sentence = nltk.sent_tokenize(f)
print("Sentences: ")
print(sentence)
words = []
words2 = []
for i in sentence:
    #word = nltk.word_tokenize(i)
    word = re.sub(r'\d+', '', i)
    words.append(word)
print("Words without numbers: ")
print(word)


def removepunc(text):
    result = ""
    for i in text:
        if i not in punctuation:
            result += i
    return result


withoutpunc = []
for i in range(len(words)):
    x = removepunc(words[i])
    print(x)
    withoutpunc.append(x)
print("Words without punctuation: ")
print(withoutpunc)

for i in withoutpunc:
    word = nltk.word_tokenize(i)
    words2.append(word)
print("Words: ")
print(words2)

onelist = []
for x in words2:
    for y in x:
        onelist.append(y)
print("onelist : ",onelist)

trigrams = []
bigrams = []
t = 0
b = 0
for i in range(len(onelist)):
    if b < i - 1:
        bigrams.append((onelist[b], onelist[b+1]))
        b += 1
print("Bigram: ", bigrams)

for i in range(len(onelist)):
    if t < i - 1:
        trigrams.append((onelist[t], onelist[t+1], onelist[t+2]))
        t += 2
print("Trigram: ", trigrams)

trigram = Counter(trigrams)
print(trigram)
countwords = Counter(onelist)
print(countwords)


def probability(a, b):
    p1 = a*b
    p2 = p1 / a


ff = len(onelist)
print(ff)
#probability()
